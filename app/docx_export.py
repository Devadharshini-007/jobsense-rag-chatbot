"""
docx_export.py
Converts plain/markdown-ish text content (from chat responses) into a
downloadable .docx file using python-docx.
"""

import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_docx_from_text(title, content_text):
    """
    Creates a .docx file in memory from a title and body text.
    Handles simple markdown-style formatting commonly produced by the LLM:
    - Lines starting with '### ' become headings
    - Lines starting with '- ' become bullet points
    - '**bold**' becomes bold text
    Returns BytesIO buffer ready for Streamlit's download_button.
    """
    doc = Document()

    # Set a clean default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title
    title_paragraph = doc.add_heading(title, level=1)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    lines = content_text.split("\n")

    for line in lines:
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph("")
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped.replace("### ", ""), level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped.replace("## ", ""), level=2)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            bullet_text = stripped[2:]
            paragraph = doc.add_paragraph(style="List Bullet")
            _add_formatted_runs(paragraph, bullet_text)
        elif stripped[0:2].isdigit() and ". " in stripped[0:4]:
            paragraph = doc.add_paragraph(style="List Number")
            number_split = stripped.split(". ", 1)
            text_part = number_split[1] if len(number_split) > 1 else stripped
            _add_formatted_runs(paragraph, text_part)
        else:
            paragraph = doc.add_paragraph()
            _add_formatted_runs(paragraph, stripped)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _add_formatted_runs(paragraph, text):
    """
    Splits text on '**bold**' markers and adds runs with appropriate
    bold formatting to the given paragraph.
    """
    parts = text.split("**")
    is_bold = False
    for part in parts:
        if part == "":
            is_bold = not is_bold
            continue
        run = paragraph.add_run(part)
        run.bold = is_bold
        is_bold = not is_bold
