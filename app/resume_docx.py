"""
resume_docx.py
Builds a properly formatted, professional-looking ATS-friendly resume .docx,
distinct from the generic docx_export.py used for fit analysis/cover letters/etc.
Resumes need real visual structure (name header, clean section dividers,
consistent spacing) that a generic markdown-to-docx converter can't provide.
"""

import io
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _add_bottom_border(paragraph):
    """Adds a clean horizontal line below a paragraph (used under section headings)."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "444444")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _parse_sections(resume_text):
    """
    Splits the LLM's '## Heading' formatted resume text into a list of
    (heading, content_lines) tuples. Content before the first heading
    (e.g. name/contact info) is returned separately as the header block.
    """
    lines = resume_text.split("\n")
    header_lines = []
    sections = []
    current_heading = None
    current_lines = []

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("## "):
            if current_heading is not None:
                sections.append((current_heading, current_lines))
            elif header_lines:
                pass
            current_heading = stripped.replace("## ", "").strip()
            current_lines = []
        elif current_heading is None:
            if stripped:
                header_lines.append(stripped)
        else:
            current_lines.append(line)

    if current_heading is not None:
        sections.append((current_heading, current_lines))

    return header_lines, sections


def _add_bold_aware_run(paragraph, text, base_size=10.5):
    """Adds text to a paragraph, handling **bold** markers within it."""
    parts = text.split("**")
    is_bold = False
    for part in parts:
        if part == "":
            is_bold = not is_bold
            continue
        run = paragraph.add_run(part)
        run.bold = is_bold
        run.font.size = Pt(base_size)
        is_bold = not is_bold


def create_resume_docx(resume_text, candidate_name="Candidate"):
    """
    Builds a polished, ATS-friendly resume .docx from the LLM's tailored
    resume text (which uses '## Heading' and '- bullet' markdown).
    Returns a BytesIO buffer ready for Streamlit's download_button.
    """
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(10.5)

    header_lines, sections = _parse_sections(resume_text)

    # Safety net: if the LLM didn't use proper '## Heading' formatting,
    # sections will be empty and we'd otherwise produce a near-blank document.
    # In that case, fall back to dumping the raw text as plain paragraphs
    # instead of silently losing all the content.
    if not sections:
        name_paragraph = doc.add_paragraph()
        name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_paragraph.add_run(candidate_name)
        name_run.bold = True
        name_run.font.size = Pt(18)
        name_paragraph.space_after = Pt(10)

        for line in resume_text.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("- ") or stripped.startswith("* "):
                bullet_paragraph = doc.add_paragraph(style="List Bullet")
                bullet_paragraph.paragraph_format.space_after = Pt(3)
                _add_bold_aware_run(bullet_paragraph, stripped[2:])
            else:
                normal_paragraph = doc.add_paragraph()
                normal_paragraph.paragraph_format.space_after = Pt(4)
                _add_bold_aware_run(normal_paragraph, stripped)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    # Name header - first line of header_lines, bold and large
    name_paragraph = doc.add_paragraph()
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_paragraph.add_run(candidate_name)
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_paragraph.space_after = Pt(2)

    # Remaining header lines (contact info, links) - centered, smaller, grey
    # Skip any line that duplicates the candidate name itself
    contact_only_lines = [
        line for line in header_lines
        if line.strip().lower() != candidate_name.strip().lower()
    ]
    if contact_only_lines:
        contact_paragraph = doc.add_paragraph()
        contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_text = " | ".join(contact_only_lines)
        contact_run = contact_paragraph.add_run(contact_text)
        contact_run.font.size = Pt(9.5)
        contact_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        contact_paragraph.space_after = Pt(10)

    # Each section: bold heading with underline rule, then content
    for heading, content_lines in sections:
        heading_paragraph = doc.add_paragraph()
        heading_paragraph.space_before = Pt(10)
        heading_paragraph.space_after = Pt(4)
        heading_run = heading_paragraph.add_run(heading.upper())
        heading_run.bold = True
        heading_run.font.size = Pt(11.5)
        heading_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        _add_bottom_border(heading_paragraph)

        for line in content_lines:
            stripped = line.strip()
            if not stripped:
                continue

            if stripped.startswith("- ") or stripped.startswith("* "):
                bullet_text = stripped[2:]
                bullet_paragraph = doc.add_paragraph(style="List Bullet")
                bullet_paragraph.paragraph_format.space_after = Pt(3)
                bullet_paragraph.paragraph_format.left_indent = Inches(0.25)
                _add_bold_aware_run(bullet_paragraph, bullet_text)
            else:
                normal_paragraph = doc.add_paragraph()
                normal_paragraph.paragraph_format.space_after = Pt(3)
                _add_bold_aware_run(normal_paragraph, stripped)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def extract_candidate_name(resume_text):
    """
    Best-effort extraction of the candidate's name from the top of their
    original resume text, used as the header in the tailored docx.
    Falls back to a generic label if nothing clear is found.
    """
    lines = [l.strip() for l in resume_text.split("\n") if l.strip()]
    for line in lines[:5]:
        if "@" in line or "http" in line.lower() or any(char.isdigit() for char in line):
            continue
        words = line.split()
        if 1 <= len(words) <= 4 and all(w[0].isupper() for w in words if w):
            return line
    return "Resume"
